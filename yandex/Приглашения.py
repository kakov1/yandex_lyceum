from docx import Document
from sys import stdin

data = list(map(str.strip, stdin))
place, time, names = data[0], data[1], data[2:]


document = Document()
for i in names:
    if names.index(i) != 0:
        page = document.add_page_break()
    document.add_heading('Приглашение', 0)
    p = document.add_paragraph(f'Уважаемая, {i}, приглашаем Вас на праздник,'
                               f' посвященный празднику 8 марта, {place}, {time}.')
document.save('test.docx')
