from docx import Document


def markdown_to_docx(text):
    document = Document()
    print(text)
    document.add_heading(text[0], 0)
    for i in text[1:]:
        if i[0] == '#':
            document.add_heading(i[i.find(' ') + 1:], level=i.find(' '))
        elif i[0] == '*' or i[0] == '-' or i[0] == '+':
            previous = text[text.index(i) - 1]
            if previous[0] == '*' or previous[0] == '-' or previous[0] == '+':
                paragraph.add_run(i[2:])
            else:
                paragraph = document.add_paragraph(i[2:], style='List Bullet')
        elif i[0].isdigit() and i[1] == '.':
            previous = text[text.index(i) - 1]
            if previous[0].isdigit() and previous[1] == '.':
                paragraph.add_run(i[2:])
            else:
                paragraph = document.add_paragraph(i[2:], style='List Number')
    document.save('res.docx')


markdown_to_docx(open('text', encoding='utf-8').readlines())